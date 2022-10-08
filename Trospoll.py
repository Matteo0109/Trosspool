from docx import Document 
from docx.enum.section import WD_ORIENT
import random
import numpy as np
from tkinter import *

def createList(n):
    lst = []
    for i in range(n+1):
        lst.append(i)
    random.shuffle(lst)
    return(lst)

  
if __name__=='__main__':
    listo = []
    #Creating the window and the frame
    main_win=Tk()
    main_win.title('Trospol')

    text_frame = Frame(main_win)
    text_frame.config(pady=10)
    text_frame.pack()

    number_frame = Frame(main_win)
    number_frame.config(background="white")
    number_frame.pack()
    
    #The entry
    e=Entry(text_frame,width=50,font = "Raleway")
    e.grid(row=0,columnspan=4)
    e.insert(0, "Mets moi un titre")

    #Fake title
    the_title = Label(text_frame, text = "Hello",font = "Raleway", pady=5)

    def create_grid():
        global listo
        listo = createList(59)
        table(listo)

    def list_to_dict():
        splits = np.array_split(listo, 5)
        dic = {"row 1":None ,"row 2":None,"row 3":None,"row 4":None,"row 5":None}
        count = 0
        for key in dic:
            dic[key]=list(splits[count])
            count +=1
        return dic

    #Print the text
    def print_text():
        global the_title
        the_title.grid_forget()
        the_title = Label(text_frame, text=e.get(), font = ("Raleway",20) )
        the_title.grid(row = 2,column =0, columnspan=4,pady=25)

    #Implementing enter key
    def enter_key(event):
        print_text()

    main_win.bind('<Return>', enter_key)

    #Real Title function
    def reset_grid():
        list_2 = []
        for x in range(60):
            list_2.append(" ")
        table(list_2)

    def create_doc():

        document= Document()

        #Changing into landscape format
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        #Doing the head
        head = e.get()
        document.add_heading(head,0)

        #Adding the table
        table = document.add_table(rows=5,cols=12, style ="Table Grid")

        #adding all element into the grid of word
        dict = list_to_dict()
        count_r=0
        for row in dict:
            data_row = table.rows[count_r].cells
            count_c=0
            for column in dict[row]:
                if count_c <13:
                    cell_value = str(column)
                    data_row[count_c].text=cell_value
                    count_c+=1
            count_r+=1

        #Saving
        document.save('Trosspool.docx')


    #Button
    entry_button = Button(text_frame, text = "Text",font= "Raleway", command=print_text, bg="#20bebe", fg= "white", height= 3, width= 15)
    entry_button.grid(row = 1, column =0, pady=15)

    create_grid_button = Button(text_frame, text = "Grille",font= "Raleway", command=create_grid, bg="#20bebe", fg= "white", height= 3, width= 15)
    create_grid_button.grid(row = 1, column =1, pady=15)

    empty_grid = Button(text_frame, text = "Reset",font= "Raleway", command=reset_grid, bg="#20bebe", fg= "white", height= 3, width= 15)
    empty_grid.grid(row = 1, column =2, pady=15)  

    save_to_docx = Button(text_frame, text = "Word",font= "Raleway",command = create_doc, bg="#20bebe", fg= "white", height= 3, width= 15)
    save_to_docx.grid(row = 1, column =3, pady=15)

    #Table
    def table(listo):
        count = 0
        for num in listo:
            if count <12 :
                lab=Label(number_frame, text=num,width=12,font = "Raleway", bg = "white",height=6, borderwidth=2, relief= "solid").grid(row=1, column=count)
                count +=1
            elif 12 <= count < 24 :
                lab=Label(number_frame, text=num,width=12,font = "Raleway", bg = "white",height=6, borderwidth=2, relief= "solid").grid(row=2, column=count-12)
                count +=1
            elif 24 <= count < 36 :
                lab=Label(number_frame, text=num,width=12,font = "Raleway", bg = "white",height=6, borderwidth=2, relief= "solid").grid(row=3, column=count-24)
                count +=1
            elif 36 <= count < 48 :
                lab=Label(number_frame, text=num,width=12,font = "Raleway", bg = "white",height=6, borderwidth=2, relief= "solid").grid(row=4, column=count-36)
                count +=1
            else:
                lab=Label(number_frame, text=num,width=12,font = "Raleway", bg = "white",height=6, borderwidth=2, relief= "solid").grid(row=5, column=count-48)
                count +=1
                
    main_win.mainloop()